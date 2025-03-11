# This file has been renamed from main.py to app.py

import os
import shutil
from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from pydantic import BaseModel
import fitz  # PyMuPDF
import docx
from langchain_community.llms import OpenAI
from langchain_core.prompts import PromptTemplate
from langchain.chains import LLMChain
from faster_whisper import WhisperModel
from TTS.api import TTS  # Correct import for Coqui TTS
from smolagents import SmolAgent
from dotenv import load_dotenv  # Import dotenv
import gradio as gr
import tempfile

load_dotenv()  # Load environment variables from .env file

app = FastAPI()

# Create temp directory if it doesn't exist
TEMP_DIR = os.path.join(os.getcwd(), "temp")
os.makedirs(TEMP_DIR, exist_ok=True)

class UserInput(BaseModel):
    user_input: str

class JobContext(BaseModel):
    job_role: str
    job_description: str

# Initialize the OpenAI LLM with your API key from the environment variable
llm = OpenAI(api_key=os.getenv('API_KEY'))  # Use the API key from .env

# Initialize models only when needed to save memory and speed up builds
whisper_model = None
tts_model = None

def get_whisper_model():
    global whisper_model
    if whisper_model is None:
        whisper_model = WhisperModel("tiny", download_root=TEMP_DIR)
    return whisper_model

def get_tts_model():
    global tts_model
    if tts_model is None:
        tts_model = TTS(model_name="tts_models/en/ljspeech/tacotron2-DDC", progress_bar=False)
    return tts_model

# Create a prompt template for adaptive questioning
prompt_template = PromptTemplate(
    input_variables=["user_input", "job_role", "job_description"],
    template="Based on the user's previous responses: {user_input}, job role: {job_role}, job description: {job_description}, generate an adaptive interview question."
)

# Create a LangChain LLM chain
chain = LLMChain(llm=llm, prompt=prompt_template)

# Initialize SmolAgent
smol_agent = SmolAgent()

@app.get("/status")
async def read_status():
    return {"status": "ok"}

@app.post("/generate-interview")
async def generate_interview(user_input: UserInput, job_context: JobContext):
    recent_context = user_input.user_input

    # Use SmolAgent to determine the next question based on the current context
    next_question = smol_agent.decide_next_question(recent_context)

    # Generate the adaptive prompt using LangChain
    adaptive_prompt = chain.run(user_input=recent_context, job_role=job_context.job_role, job_description=job_context.job_description)

    # Combine the SmolAgent's decision with the adaptive prompt
    generated_text = f"{next_question} {adaptive_prompt}"
    return {"generated_text": generated_text}

def transcribe_audio(audio_path: str) -> str:
    # Use Whisper to transcribe audio to text
    model = get_whisper_model()
    segments, _ = model.transcribe(audio_path)
    transcribed_text = " ".join([segment.text for segment in segments])
    return transcribed_text

def synthesize_speech(text: str, output_path: str):
    # Use Coqui TTS to convert text to speech
    model = get_tts_model()
    model.tts_to_file(text=text, file_path=output_path)

@app.post("/process-audio")
async def process_audio(audio: UploadFile = File(...)):
    # Create a temporary file with a proper extension
    with tempfile.NamedTemporaryFile(delete=False, suffix='.wav', dir=TEMP_DIR) as temp_audio:
        audio_path = temp_audio.name
        shutil.copyfileobj(audio.file, temp_audio)

    try:
        # Transcribe the audio to text
        transcribed_text = transcribe_audio(audio_path)

        # Process the text response through the existing Gemini API/LangChain adaptive logic
        job_context = JobContext(job_role="example_role", job_description="example_description")
        adaptive_prompt = chain.run(user_input=transcribed_text, job_role=job_context.job_role, job_description=job_context.job_description)

        # Synthesize speech from the generated feedback
        output_audio_path = os.path.join(TEMP_DIR, 'feedback_audio.wav')
        synthesize_speech(adaptive_prompt, output_audio_path)

        return {"audio_url": output_audio_path}
    finally:
        # Clean up temporary files
        if os.path.exists(audio_path):
            os.unlink(audio_path)

@app.post("/submit-context")
async def submit_context(job_role: str = Form(...), job_description: str = Form(...), resume: UploadFile = File(...)):
    if resume.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Invalid file type. Only PDF and DOCX are allowed.")

    resume_text = ""
    if resume.content_type == "application/pdf":
        pdf_document = fitz.open(stream=await resume.read(), filetype="pdf")
        resume_text = " ".join([page.get_text() for page in pdf_document])
    elif resume.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(resume.file)
        resume_text = " ".join([para.text for para in doc.paragraphs])

    return {"job_role": job_role, "job_description": job_description, "resume_text": resume_text}

# Gradio handler functions
def gradio_process_audio(audio_data):
    if audio_data is None:
        return None
    
    # Extract audio data and save to a temporary file
    audio_path = os.path.join(TEMP_DIR, "temp_audio.wav")
    sr, y = audio_data
    import scipy.io.wavfile as wav
    wav.write(audio_path, sr, y)
    
    try:
        # Transcribe the audio to text
        transcribed_text = transcribe_audio(audio_path)
        
        # Process the text response
        job_context = JobContext(job_role="example_role", job_description="example_description")
        adaptive_prompt = chain.run(user_input=transcribed_text, job_role=job_context.job_role, job_description=job_context.job_description)
        
        # Synthesize speech from the generated feedback
        output_audio_path = os.path.join(TEMP_DIR, 'feedback_audio.wav')
        synthesize_speech(adaptive_prompt, output_audio_path)
        
        return output_audio_path
    except Exception as e:
        print(f"Error processing audio: {e}")
        return None
    finally:
        # Clean up
        if os.path.exists(audio_path):
            os.unlink(audio_path)

def gradio_submit_context(job_role, job_description, resume_file):
    if not resume_file:
        return "Please upload a resume file."
    
    try:
        resume_text = ""
        file_path = resume_file.name
        
        if file_path.endswith('.pdf'):
            pdf_document = fitz.open(file_path)
            resume_text = " ".join([page.get_text() for page in pdf_document])
        elif file_path.endswith('.docx'):
            doc = docx.Document(file_path)
            resume_text = " ".join([para.text for para in doc.paragraphs])
        else:
            return "Invalid file type. Only PDF and DOCX are allowed."
        
        return f"Job Role: {job_role}\nJob Description: {job_description}\nResume Text: {resume_text[:500]}..."
    except Exception as e:
        return f"Error processing resume: {e}"

# Create a single Gradio interface that combines all functionality
def create_gradio_interface():
    with gr.Blocks(title="AI Interviewer") as interface:
        with gr.Tab("Audio Processing"):
            audio_input = gr.Audio()
            audio_output = gr.Audio()
            audio_button = gr.Button("Process Audio")
            audio_button.click(fn=gradio_process_audio, inputs=audio_input, outputs=audio_output)
            
        with gr.Tab("Submit Context"):
            job_role = gr.Textbox(label="Job Role")
            job_desc = gr.Textbox(label="Job Description")
            resume = gr.File(label="Resume")
            submit_button = gr.Button("Submit")
            output = gr.Textbox(label="Result")
            submit_button.click(fn=gradio_submit_context, inputs=[job_role, job_desc, resume], outputs=output)
    
    return interface

# Launch the application
if __name__ == "__main__":
    import uvicorn
    interface = create_gradio_interface()
    interface.launch()
    uvicorn.run(app, host="0.0.0.0", port=7860)
